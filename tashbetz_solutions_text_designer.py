import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """מגדיר את הפסקה כ-RTL"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_run_rtl_and_format(run, is_bold=False, is_underline=False, font_size=None):
    """
    פונקציה חכמה שמחילה גם RTL וגם עיצובים (Bold/Underline)
    בצורה שתתאים ספציפית לשפה העברית.
    """
    r = run._element
    rPr = r.get_or_add_rPr()
    
    # 1. הגדרת הטקסט כעברית (RTL)
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    
    # 2. טיפול בהדגשה (BOLD)
    if is_bold:
        run.bold = True # מדגיש מספרים/אנגלית
        
        # הפקודה הקריטית לעברית: Complex Script Bold
        bCs = OxmlElement('w:bCs') 
        bCs.set(qn('w:val'), '1')
        rPr.append(bCs)
        
    # 3. טיפול בקו תחתון
    if is_underline:
        run.underline = True
        
    # 4. שינוי גודל גופן במידת הצורך
    if font_size:
        run.font.size = font_size

def export_crossword_to_rich_text(excel_path, output_docx_path):
    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        print(f"שגיאה בקריאת קובץ האקסל: {e}")
        return

    # סידור תאריכים
    if pd.api.types.is_datetime64_any_dtype(df['תאריך התשבץ']):
        df['תאריך התשבץ'] = df['תאריך התשבץ'].dt.strftime('%d/%m/%Y')
    else:
        df['תאריך התשבץ'] = df['תאריך התשבץ'].astype(str)

    doc = Document()
    
    # הגדרת סגנון גופן נקי שקורא עברית טוב
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial' 
    font.size = Pt(12)

    grouped = df.groupby(['כותרת התשבץ', 'תאריך התשבץ'], sort=False)

    for (title, date), group in grouped:
        # --- כותרת התשבץ ---
        p_header = doc.add_paragraph()
        set_rtl(p_header) 
        
        run_header = p_header.add_run(f"תשבץ מס' {title} ({date}):")
        # קריאה לפונקציה שתעשה גם RTL, גם מודגש וגם קו תחתון
        set_run_rtl_and_format(run_header, is_bold=True, is_underline=True, font_size=Pt(14))

        # --- הגדרות ---
        for _, row in group.iterrows():
            clue_num = str(row['מספר הגדרה']) if pd.notna(row['מספר הגדרה']) else ""
            solution = str(row['פתרון']) if pd.notna(row['פתרון']) else ""
            explanation = str(row['הסבר']) if pd.notna(row['הסבר']) else ""

            p = doc.add_paragraph()
            set_rtl(p)
            
            # מספר הגדרה: מודגש + קו תחתון
            run_num = p.add_run(clue_num)
            set_run_rtl_and_format(run_num, is_bold=True, is_underline=True)
            
            # פתרון: טקסט רגיל
            run_sol = p.add_run(f" - {solution}. ")
            set_run_rtl_and_format(run_sol)
            
            # "הסבר:": מודגש בלבד
            run_exp_label = p.add_run("הסבר: ")
            set_run_rtl_and_format(run_exp_label, is_bold=True)
            
            # תוכן ההסבר: טקסט רגיל
            run_exp_val = p.add_run(explanation)
            set_run_rtl_and_format(run_exp_val)

        # רווח בין תשבצים
        doc.add_paragraph()

    doc.save(output_docx_path)
    print(f"הקובץ נשמר בהצלחה כולל Bold ו-RTL: {output_docx_path}")

if __name__ == "__main__":
    # עדכון שם הקובץ לשם שביקשת
    input_excel = 'crosswords_hybrid_20250808_to_20260313 (3).xlsx'
    output_file = 'crosswords_formatted_output.docx'
    
    export_crossword_to_rich_text(input_excel, output_file)